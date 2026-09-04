from fastapi import FastAPI, APIRouter, File, UploadFile, HTTPException, Header, Depends
from starlette.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from openai import AsyncOpenAI
import groq
import httpx
import json
import math
import os
import logging
import time
from contextlib import asynccontextmanager
from pathlib import Path
from collections import defaultdict
from pydantic import BaseModel, Field
from typing import Any, Dict, List, Optional


# ── OpenTelemetry Setup ────────────────────────────────────────────────────────
# Instruments FastAPI (HTTP spans), OpenAI (chat completions, TTS, Whisper),
# and Groq (chat completions, Whisper). Traces export to Phoenix (OTLP HTTP).
# Run `phoenix server` locally, or set PHOENIX_ENDPOINT / OTEL_EXPORTER_OTLP_ENDPOINT.
# Set OTEL_SDK_DISABLED=true to disable without touching code.
def _setup_telemetry(app):
    """Set up OpenTelemetry tracing and instrument AI SDKs. Idempotent."""
    try:
        from opentelemetry import trace
        from opentelemetry.sdk.trace import TracerProvider
        from opentelemetry.sdk.resources import Resource
        from opentelemetry.sdk.trace.export import BatchSpanProcessor
        from opentelemetry.exporter.otlp.proto.http.trace_exporter import OTLPSpanExporter
        from opentelemetry.instrumentation.fastapi import FastAPIInstrumentor
        from opentelemetry.semconv.resource import ResourceAttributes
    except ImportError:
        logging.getLogger(__name__).warning(
            "OpenTelemetry packages missing — tracing disabled. Run:\n"
            "  pip install opentelemetry-sdk opentelemetry-exporter-otlp-proto-http \\\n"
            "    opentelemetry-instrumentation-fastapi \\\n"
            "    opentelemetry-instrumentation-openai-v2\n"
        )
        return

    if os.environ.get("OTEL_SDK_DISABLED", "").lower() in ("1", "true", "yes"):
        return

    resource = Resource(attributes={
        ResourceAttributes.SERVICE_NAME: "credeasy-backend",
        ResourceAttributes.SERVICE_VERSION: "1.0.0",
    })
    provider = TracerProvider(resource=resource)
    trace.set_tracer_provider(provider)

    # Export to Phoenix OTLP endpoint
    otel_endpoint = os.environ.get(
        "OTEL_EXPORTER_OTLP_ENDPOINT",
        os.environ.get("PHOENIX_ENDPOINT", "http://localhost:6006")
    ).rstrip("/")

    try:
        exporter = OTLPSpanExporter(endpoint=f"{otel_endpoint}/v1/traces")
        provider.add_span_processor(BatchSpanProcessor(exporter))
        logging.getLogger(__name__).info(
            f"OpenTelemetry active — traces → {otel_endpoint}"
        )
    except Exception as e:
        logging.getLogger(__name__).warning(f"Failed to configure OTLP exporter: {e}")

    try:
        FastAPIInstrumentor.instrument_app(app)
    except Exception as e:
        logging.getLogger(__name__).warning(f"Failed to instrument FastAPI: {e}")

    # Capture prompt/completion content (toggle via OTEL_INSTRUMENTATION_GENAI_CAPTURE_MESSAGE_CONTENT)
    os.environ.setdefault(
        "OTEL_INSTRUMENTATION_GENAI_CAPTURE_MESSAGE_CONTENT", "span_and_event"
    )
    try:
        from opentelemetry.instrumentation.openai_v2 import OpenAIInstrumentor
        OpenAIInstrumentor().instrument()
    except Exception as e:
        logging.getLogger(__name__).warning(f"Failed to instrument OpenAI SDK: {e}")


# Configured before anything else so import-time failures below are logged
# rather than raising NameError on `logger`.
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Supabase issues and signs the app's session tokens, so this server verifies
# them rather than minting its own. The anon key is public by design — it is
# sent as the `apikey` header Supabase's auth API requires, not as a secret.
SUPABASE_URL = os.environ.get("SUPABASE_URL", "").rstrip("/")
SUPABASE_ANON_KEY = os.environ.get("SUPABASE_ANON_KEY", "")


@asynccontextmanager
async def lifespan(_app: FastAPI):
    yield


app = FastAPI(lifespan=lifespan)

# Wire up OpenTelemetry after `app` exists. Idempotent and safe to call with
# OTel packages missing.
_setup_telemetry(app)

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Add your routes to the router instead of directly to app
VOICE_RATE_LIMIT_PER_MINUTE = 20
VOICE_RATE_LIMIT_WINDOW_SECONDS = 60
USER_RATE_LIMIT_BUCKETS = defaultdict(list)
# Buckets are swept only once the dict grows past this, so the common path stays
# O(1). Without any eviction it gained one list per user for the process's life.
MAX_RATE_LIMIT_BUCKETS = 10_000

UPSTREAM_TIMEOUT_SECONDS = 10.0
MAX_AUDIO_BYTES = 25 * 1024 * 1024
MAX_TRANSCRIPT_LEN = 2000
MAX_CONTEXT_CHARS = 64 * 1024
MAX_REPLY_LEN = 4000
MAX_ACTIONS = 6
MAX_NAME_LEN = 100
MAX_NOTE_LEN = 200
MAX_PHONE_LEN = 20
ALLOWED_ROUTES = {"dashboard", "parties", "billing", "reports", "settings"}


class VoiceAssistRequest(BaseModel):
    transcript: str = Field(default="", max_length=MAX_TRANSCRIPT_LEN)
    context: Dict[str, Any] = Field(default_factory=dict)
    lang: str = Field(default="en", max_length=8)
    history: List[Dict[str, Any]] = Field(default_factory=list)


class SpeakRequest(BaseModel):
    text: str = Field(default="", max_length=2000)
    lang: str = Field(default="en", max_length=8)


async def get_authenticated_user(authorization: Optional[str] = Header(None)) -> dict:
    """Resolve the current user from a Supabase access token. Raises 401, never 403.

    This replaces a second, parallel session system: the server used to mint its
    own `st_...` tokens into db.user_sessions, while the app only ever held a
    Supabase JWT. The two could never agree, so every authenticated route
    rejected every request. Asking Supabase to validate its own token leaves one
    source of truth and lets the Mongo users/user_sessions collections go away.
    """
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid Authorization header")
    token = authorization.split(" ", 1)[1].strip()
    if not token:
        raise HTTPException(status_code=401, detail="Missing session token")

    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        # A misconfigured server must not masquerade as a rejected user, or the
        # whole fleet looks like every shopkeeper's token expired at once.
        logger.error("SUPABASE_URL / SUPABASE_ANON_KEY are not configured")
        raise HTTPException(status_code=500, detail="Authentication is not configured on the server")

    try:
        async with httpx.AsyncClient(timeout=UPSTREAM_TIMEOUT_SECONDS) as httpx_client:
            resp = await httpx_client.get(
                f"{SUPABASE_URL}/auth/v1/user",
                headers={"Authorization": f"Bearer {token}", "apikey": SUPABASE_ANON_KEY},
            )
    except Exception:
        # 503, not 401: "we could not check" is not "your token is bad", and
        # returning 401 for a social login blip signs the user out of the app.
        logger.exception("Supabase token verification request failed")
        raise HTTPException(status_code=503, detail="Could not verify your session. Please try again.")

    if resp.status_code != 200:
        raise HTTPException(status_code=401, detail="Invalid or expired session")

    try:
        data = resp.json()
    except ValueError:
        logger.error("Supabase /auth/v1/user returned a non-JSON body")
        raise HTTPException(status_code=502, detail="Could not verify your session. Please try again.")

    if not isinstance(data, dict):
        raise HTTPException(status_code=502, detail="Could not verify your session. Please try again.")

    user_id = data.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid or expired session")

    metadata = data.get("user_metadata")
    if not isinstance(metadata, dict):
        metadata = {}

    return {
        "user_id": user_id,
        "email": data.get("email") or metadata.get("email") or "",
        "name": metadata.get("full_name") or data.get("confirmed_at") or user_id[:8],
    }


async def get_optional_user(authorization: Optional[str] = Header(None)) -> Optional[dict]:
    """Resolve the current user if a valid token is provided; return None if absent.

    Use this for endpoints where authentication is helpful but not mandatory — e.g.
    file parsing during onboarding where the user may not have a persistent session.
    """
    if not authorization or not authorization.startswith("Bearer "):
        return None
    token = authorization.split(" ", 1)[1].strip()
    if not token:
        return None
    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        return None

    try:
        async with httpx.AsyncClient(timeout=UPSTREAM_TIMEOUT_SECONDS) as httpx_client:
            resp = await httpx_client.get(
                f"{SUPABASE_URL}/auth/v1/user",
                headers={"Authorization": f"Bearer {token}", "apikey": SUPABASE_ANON_KEY},
            )
    except Exception:
        return None

    if resp.status_code != 200:
        return None

    try:
        data = resp.json()
    except ValueError:
        return None

    if not isinstance(data, dict) or not data.get("id"):
        return None

    metadata = data.get("user_metadata") or {}
    return {
        "user_id": data["id"],
        "email": data.get("email") or metadata.get("email") or "",
        "name": metadata.get("full_name") or data.get("confirmed_at") or data["id"][:8],
    }
    """Resolve the current user from a Supabase access token. Raises 401, never 403.

    This replaces a second, parallel session system: the server used to mint its
    own `st_...` tokens into db.user_sessions, while the app only ever held a
    Supabase JWT. The two could never agree, so every authenticated route
    rejected every request. Asking Supabase to validate its own token leaves one
    source of truth and lets the Mongo users/user_sessions collections go away.
    """
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid Authorization header")
    token = authorization.split(" ", 1)[1].strip()
    if not token:
        raise HTTPException(status_code=401, detail="Missing session token")

    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        # A misconfigured server must not masquerade as a rejected user, or the
        # whole fleet looks like every shopkeeper's token expired at once.
        logger.error("SUPABASE_URL / SUPABASE_ANON_KEY are not configured")
        raise HTTPException(status_code=500, detail="Authentication is not configured on the server")

    try:
        async with httpx.AsyncClient(timeout=UPSTREAM_TIMEOUT_SECONDS) as httpx_client:
            resp = await httpx_client.get(
                f"{SUPABASE_URL}/auth/v1/user",
                headers={"Authorization": f"Bearer {token}", "apikey": SUPABASE_ANON_KEY},
            )
    except Exception:
        # 503, not 401: "we could not check" is not "your token is bad", and
        # returning 401 for a network blip signs the user out of the app.
        logger.exception("Supabase token verification request failed")
        raise HTTPException(status_code=503, detail="Could not verify your session. Please try again.")

    if resp.status_code != 200:
        raise HTTPException(status_code=401, detail="Invalid or expired session")

    try:
        data = resp.json()
    except ValueError:
        logger.error("Supabase /auth/v1/user returned a non-JSON body")
        raise HTTPException(status_code=502, detail="Could not verify your session. Please try again.")

    if not isinstance(data, dict):
        raise HTTPException(status_code=502, detail="Could not verify your session. Please try again.")

    user_id = data.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid or expired session")

    metadata = data.get("user_metadata")
    if not isinstance(metadata, dict):
        metadata = {}

    return {
        "user_id": str(user_id),
        "email": data.get("email") or "",
        "name": metadata.get("full_name") or metadata.get("name") or "",
        "picture": metadata.get("avatar_url"),
    }


def _evict_stale_rate_limit_buckets(now: float, window_seconds: int) -> None:
    if len(USER_RATE_LIMIT_BUCKETS) <= MAX_RATE_LIMIT_BUCKETS:
        return
    stale = [
        user_id for user_id, timestamps in USER_RATE_LIMIT_BUCKETS.items()
        if not timestamps or now - timestamps[-1] >= window_seconds
    ]
    for user_id in stale:
        del USER_RATE_LIMIT_BUCKETS[user_id]


def enforce_user_rate_limit(user_id: str, limit: int = VOICE_RATE_LIMIT_PER_MINUTE, window_seconds: int = VOICE_RATE_LIMIT_WINDOW_SECONDS):
    now = time.monotonic()
    _evict_stale_rate_limit_buckets(now, window_seconds)
    timestamps = USER_RATE_LIMIT_BUCKETS[user_id]
    timestamps[:] = [ts for ts in timestamps if now - ts < window_seconds]
    if len(timestamps) >= limit:
        raise HTTPException(status_code=429, detail="Rate limit exceeded. Please try again shortly.")
    timestamps.append(now)


@api_router.get("/auth/me")
async def auth_me(user: dict = Depends(get_authenticated_user)):
    return {
        "user": {
            "user_id": user["user_id"],
            "email": user.get("email"),
            "name": user.get("name"),
            "picture": user.get("picture"),
        }
    }


@api_router.delete("/auth/account")
async def delete_account(user: dict = Depends(get_authenticated_user)):
    """Delete the authenticated user's account from Supabase Auth.

    This endpoint requires the SUPABASE_SERVICE_ROLE_KEY environment variable.
    Without it, this endpoint returns 500.
    """
    service_role_key = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
    if not service_role_key:
        logger.error("SUPABASE_SERVICE_ROLE_KEY is not configured")
        raise HTTPException(status_code=500, detail="Account deletion is not configured on the server")

    user_id = user["user_id"]

    try:
        async with httpx.AsyncClient(timeout=UPSTREAM_TIMEOUT_SECONDS) as httpx_client:
            resp = await httpx_client.delete(
                f"{SUPABASE_URL}/auth/v1/admin/users/{user_id}",
                headers={
                    "Authorization": f"Bearer {service_role_key}",
                    "apikey": service_role_key,
                    "Content-Type": "application/json"
                },
            )

        if resp.status_code in (200, 204):
            return {"success": True, "message": "Account deleted successfully"}
        else:
            logger.error(f"Failed to delete account: {resp.status_code} {resp.text}")
            raise HTTPException(status_code=500, detail="Failed to delete account. Please try again.")

    except Exception as e:
        logger.exception("Account deletion request failed")
        raise HTTPException(status_code=500, detail="Could not delete account. Please try again.")


@api_router.get("/")
async def root():
    return {"message": "Hello World"}


@api_router.get("/health")
async def health():
    """Liveness probe — reports service status and dependency readiness.

    Returns 200 with a JSON body describing each external dependency the server
    relies on. Useful for uptime monitoring and pre-deploy checks. Does not
    require auth so external monitors can hit it without a Supabase token.
    """
    deps = {
        "supabase": bool(SUPABASE_URL and SUPABASE_ANON_KEY),
        "groq": bool(os.environ.get("GROQ_API_KEY") or os.environ.get("GROQ_APIKEY")),
        "openai": bool(
            os.environ.get("OPENAI_API_KEY")
            or os.environ.get("LLM_API_KEY")
            or os.environ.get("EMERGENT_LLM_KEY")
        ),
    }
    return {
        "status": "ok",
        "service": "credeasy-backend",
        "dependencies": deps,
    }


# ---------------------------------------------------------------------------
# Voice Assistant: Whisper transcription + LLM intent parsing
# ---------------------------------------------------------------------------
ALLOWED_AUDIO_EXT = {".m4a", ".mp3", ".mp4", ".mpeg", ".mpga", ".wav", ".webm"}

_openai_client: Optional[AsyncOpenAI] = None
_groq_client: Optional[groq.AsyncGroq] = None


def get_groq_api_key() -> Optional[str]:
    """Groq is the primary provider — free tier, no credit card needed.
    Used for chat completions and Whisper transcription."""
    key = os.environ.get("GROQ_API_KEY") or os.environ.get("GROQ_APIKEY") or None
    if key:
        logger.info(f"GROQ_API_KEY found (length={len(key)})")
    else:
        logger.warning("GROQ_API_KEY not found in environment")
    return key


def get_openai_api_key() -> Optional[str]:
    """OpenAI is the fallback if Groq is unavailable or exhausted."""
    return (
        os.environ.get("OPENAI_API_KEY")
        or os.environ.get("LLM_API_KEY")
        or os.environ.get("EMERGENT_LLM_KEY")
        or None
    )


def get_groq_client() -> groq.AsyncGroq:
    global _groq_client
    if _groq_client is None:
        api_key = get_groq_api_key()
        if not api_key:
            raise HTTPException(status_code=500, detail="GROQ_API_KEY is not configured")
        _groq_client = groq.AsyncGroq(api_key=api_key)
    return _groq_client


def get_openai_client() -> AsyncOpenAI:
    global _openai_client
    if _openai_client is None:
        api_key = get_openai_api_key()
        if not api_key:
            raise HTTPException(status_code=500, detail="OPENAI_API_KEY is not configured")
        _openai_client = AsyncOpenAI(api_key=api_key)
    return _openai_client


@api_router.post("/voice/transcribe")
async def voice_transcribe(file: UploadFile = File(...), user: dict = Depends(get_authenticated_user)):
    enforce_user_rate_limit(str(user["user_id"]))

    suffix = Path(file.filename or "recording.m4a").suffix.lower() or ".m4a"
    if suffix not in ALLOWED_AUDIO_EXT:
        raise HTTPException(status_code=415, detail=f"Unsupported audio format: {suffix}")

    audio = await file.read()
    if not audio:
        raise HTTPException(status_code=400, detail="Empty audio file")
    if len(audio) > MAX_AUDIO_BYTES:
        raise HTTPException(status_code=413, detail="Audio must be 25 MB or smaller")

    # Try Groq first (free Whisper), fall back to OpenAI.
    last_error = None
    if get_groq_api_key():
        try:
            groq_client = get_groq_client()
            result = await groq_client.audio.transcriptions.create(
                model="whisper-large-v3-turbo",  # faster, free tier model
                file=(f"recording{suffix}", audio),
                prompt="Hindi and English (Hinglish) shopkeeper ledger speech. Preserve names, numbers and rupee amounts.",
            )
            return {"text": getattr(result, "text", "") or ""}
        except Exception as e:
            logger.warning("Groq transcription failed, falling back to OpenAI: %s", e)
            last_error = e

    if get_openai_api_key():
        try:
            openai_client = get_openai_client()
            result = await openai_client.audio.transcriptions.create(
                model="whisper-1",
                file=(f"recording{suffix}", audio),
                prompt="Hindi and English (Hinglish) shopkeeper ledger speech. Preserve names, numbers and rupee amounts.",
            )
            return {"text": getattr(result, "text", "") or ""}
        except Exception as e:
            logger.exception("OpenAI transcription failed")
            last_error = e

    logger.exception("transcription failed: no provider succeeded")
    raise HTTPException(status_code=502, detail="Could not transcribe the recording. Please try again.")


@api_router.post("/voice/speak")
async def voice_speak(payload: SpeakRequest, user: dict = Depends(get_authenticated_user)):
    """Synthesize text to speech and stream the MP3 back.

    Uses Edge TTS (free, no API key) for TTS, with OpenAI as fallback.
    Returns audio as a streaming response so the client can play() while
    bytes are still arriving. If TTS fails for any reason, the assistant's
    on-screen text reply still works — speech is a progressive enhancement.
    """
    enforce_user_rate_limit(str(user["user_id"]))

    text = (payload.text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="text is required")
    text = text[:2000]

    # Pick a voice that handles Hindi + English well. Edge TTS voices are free.
    lang = (payload.lang or "en").lower()
    # Hindi voices: hi-IN-SwaraNeural (female), hi-IN-MadhurNeural (male)
    # English voices: en-US-JennyNeural, en-US-GuyNeural
    if lang.startswith("hi"):
        voice = "hi-IN-SwaraNeural"
    else:
        voice = "en-US-JennyNeural"

    # Try Edge TTS first (completely free, no API key needed)
    try:
        import edge_tts

        async def audio_generator():
            """Stream audio chunks as they arrive from Edge TTS."""
            communicate = edge_tts.Communicate(text, voice)
            async for chunk in communicate.stream():
                if chunk["type"] == "audio":
                    yield chunk["data"]

        return StreamingResponse(
            audio_generator(),
            media_type="audio/mpeg",
            headers={
                "Content-Disposition": "inline",
                "Accept-Ranges": "none",
            },
        )
    except ImportError:
        logger.warning("edge-tts not installed, trying fallback")
    except Exception as e:
        logger.warning("Edge TTS failed, trying OpenAI fallback: %s", e)

    # Fallback to OpenAI if available
    if get_openai_api_key():
        try:
            model = os.environ.get("TTS_MODEL", "gpt-4o-mini-tts")
            openai_client = get_openai_client()
            response = await openai_client.audio.speech.with_streaming_response.create(
                model=model,
                voice="alloy",
                input=text,
                response_format="mp3",
            )
            return StreamingResponse(
                response.iter_bytes(),
                media_type="audio/mpeg",
                headers={"Content-Disposition": "inline"},
            )
        except Exception:
            logger.exception("OpenAI TTS failed")

    raise HTTPException(status_code=502, detail="Could not synthesize speech. Please try again.")


VOICE_SYSTEM_PROMPT = """You are CredEasy Assistant, a voice helper inside CredEasy — a digital khata (credit ledger) app used by small Indian shopkeepers.
The user speaks Hindi, English or Hinglish. You do three things:
1. RECORD entries: e.g. "Ramesh ko 500 rupaye udhaar diye" -> add a GAVE transaction of 500 for party Ramesh.
2. ANSWER questions about the ledger using the CONTEXT provided: e.g. "Ramesh ka kitna baaki hai?".
3. GUIDE / navigate the user through the app: e.g. "bill kaise banau" -> explain briefly and navigate to billing.
4. ADD parties interactively: the user can add a new party step by step.

Reply ONLY with a single JSON object, no markdown, no code fences:
{
  "reply": "<short spoken-style answer, max 2 sentences, in the SAME language the user used>",
  "actions": [ ... ]
}

Allowed actions (0 to 3 items):
{"type":"ADD_TRANSACTION","partyName":"<name as spoken>","amount":<number>,"txType":"GAVE"|"GOT","note":"<short note or empty>"}
   - GAVE = shopkeeper gave goods/credit (money receivable). GOT = shopkeeper received payment.
{"type":"ASK_PARTY_SPELLING","name":"<name as spoken by user>"}
   - Ask the user to spell the party name letter by letter. Show in reply.
{"type":"SELECT_CONTACT","name":"<spelled name>"}
   - After spelling is confirmed, ask if the contact is in their phone's contact list.
{"type":"ASK_OPENING_BALANCE","name":"<name>","phone":"<phone or empty>"}
   - After contact is selected (or skipped), ask if there's an opening balance.
{"type":"ADD_PARTY_COMPLETE","name":"<name>","phone":"<phone>","openingBalance":<number>,"partyType":"CUSTOMER"|"SUPPLIER"}
   - When user confirms the amount (or says "nahi"), create the party.
{"type":"NAVIGATE","route":"dashboard"|"parties"|"billing"|"reports"|"settings"}
{"type":"REMIND","partyName":"<name>"}

## Multi-step ADD PARTY flow:
Step 1 — User says "add Ramesh" or "Ramesh ko add karo":
  → reply: "Ramesh ka spelling batao, ek ek letter bolo." (Hindi)
  → action: ASK_PARTY_SPELLING with the name

Step 2 — User spells "R A M E S H":
  → reply: "Ramesh. Kya Ramesh ka phone aapke contact list mein hai?" (Hindi)
  → action: SELECT_CONTACT with name="RAMESH"

Step 3 — Frontend sends "yes" or "no" or "haan" / "nahi":
  If yes: frontend shows matching contact cards. User taps one.
    Frontend sends "Selected: Ramesh, 9876543210"
    → reply: "9876543210. Kya koi opening balance hai? Amount bolo ya 'nahi' bolo." (Hindi)
    → action: ASK_OPENING_BALANCE with name and phone
  If no:
    → reply: "Kya koi opening balance hai? Amount bolo ya 'nahi' bolo." (Hindi)
    → action: ASK_OPENING_BALANCE with name, phone empty

Step 4 — User says amount or "nahi":
  → reply: "Sab theek hai. Ramesh ka khata ban gaya." (Hindi)
  → action: ADD_PARTY_COMPLETE with name, phone, openingBalance (0 if nahi), partyType (infer from context, default CUSTOMER)

## Rules:
- Only emit ADD_TRANSACTION when an amount AND a party are clearly stated. Otherwise ask for what is missing.
- If the user says something vague ("do something for Ramesh", "update that", "batao"), ask a clarifying question in the SAME language before acting. Do not guess.
  Example: user says "Ramesh ko 500 diye" but no Ramesh exists → ask "Kaunsa Ramesh? Ramesh Kumar ya Ramesh Sharma?"
  Example: user says "uska baaki check karo" → ask "Kaunsa customer ka baaki?"
- Match partyName to the closest existing party name from CONTEXT when possible.
- For pure questions, answer with numbers from CONTEXT and return an empty actions array.
- Amounts are Indian rupees; convert words like "paanch sau" to 500.
- Never invent balances that are not in CONTEXT.
- Keep replies short — 1-2 sentences max. The user is listening, not reading.
"""


def _clean_str(value: object, max_len: int) -> str:
    if not isinstance(value, str):
        value = "" if value is None else str(value)
    return value.strip()[:max_len]


def _clean_amount(value: object) -> Optional[float]:
    """A positive, finite rupee amount, or None. Rejects NaN, infinity and
    "1e999" — any of which would poison every balance the party appears in."""
    if isinstance(value, bool):
        return None
    try:
        amount = float(value)  # type: ignore[arg-type]
    except (TypeError, ValueError):
        return None
    if not math.isfinite(amount) or amount <= 0:
        return None
    return round(amount, 2)


def sanitize_actions(raw: object) -> List[dict]:
    """The model's output is untrusted input: the app writes these actions
    straight into the shopkeeper's ledger. Emit only whitelisted shapes with
    coerced, length-capped fields and drop anything that does not fit."""
    if not isinstance(raw, list):
        return []
    actions: List[dict] = []
    for item in raw[:MAX_ACTIONS]:
        if not isinstance(item, dict):
            continue
        kind = item.get("type")
        if kind == "ADD_TRANSACTION":
            amount = _clean_amount(item.get("amount"))
            party = _clean_str(item.get("partyName"), MAX_NAME_LEN)
            tx_type = item.get("txType")
            if amount is None or not party or tx_type not in ("GAVE", "GOT"):
                continue
            actions.append({
                "type": kind,
                "partyName": party,
                "amount": amount,
                "txType": tx_type,
                "note": _clean_str(item.get("note"), MAX_NOTE_LEN),
            })
        elif kind == "ADD_PARTY":
            name = _clean_str(item.get("name"), MAX_NAME_LEN)
            if not name:
                continue
            actions.append({
                "type": kind,
                "name": name,
                "phone": _clean_str(item.get("phone"), MAX_PHONE_LEN),
                "partyType": "SUPPLIER" if item.get("partyType") == "SUPPLIER" else "CUSTOMER",
            })
        elif kind == "ASK_PARTY_SPELLING":
            name = _clean_str(item.get("name"), MAX_NAME_LEN)
            if not name:
                continue
            actions.append({"type": kind, "name": name})
        elif kind == "SELECT_CONTACT":
            name = _clean_str(item.get("name"), MAX_NAME_LEN)
            if not name:
                continue
            actions.append({"type": kind, "name": name})
        elif kind == "ASK_OPENING_BALANCE":
            name = _clean_str(item.get("name"), MAX_NAME_LEN)
            phone = _clean_str(item.get("phone"), MAX_PHONE_LEN)
            if not name:
                continue
            actions.append({"type": kind, "name": name, "phone": phone})
        elif kind == "ADD_PARTY_COMPLETE":
            name = _clean_str(item.get("name"), MAX_NAME_LEN)
            if not name:
                continue
            balance = _clean_amount(item.get("openingBalance"))
            actions.append({
                "type": kind,
                "name": name,
                "phone": _clean_str(item.get("phone"), MAX_PHONE_LEN),
                "openingBalance": balance if balance is not None else 0,
                "partyType": "SUPPLIER" if item.get("partyType") == "SUPPLIER" else "CUSTOMER",
            })
        elif kind == "NAVIGATE":
            route = _clean_str(item.get("route"), 32)
            if route not in ALLOWED_ROUTES:
                continue
            actions.append({"type": kind, "route": route})
        elif kind == "REMIND":
            party = _clean_str(item.get("partyName"), MAX_NAME_LEN)
            if not party:
                continue
            actions.append({"type": kind, "partyName": party})
    return actions


@api_router.post("/voice/assist")
async def voice_assist(payload: VoiceAssistRequest, user: dict = Depends(get_authenticated_user)):
    enforce_user_rate_limit(str(user["user_id"]))

    transcript = payload.transcript.strip()
    if not transcript:
        raise HTTPException(status_code=400, detail="transcript is required")

    # The ledger context is pasted verbatim into the prompt, so an oversized one
    # is a token-cost amplifier as much as a memory concern.
    context_json = json.dumps(payload.context, ensure_ascii=False, default=str)
    if len(context_json) > MAX_CONTEXT_CHARS:
        raise HTTPException(status_code=413, detail="Ledger context is too large to send.")

    user_text = f"CONTEXT (current ledger):\n{context_json}\n\nUSER SAID: {transcript}"

    messages: List[Dict[str, str]] = [{"role": "system", "content": VOICE_SYSTEM_PROMPT}]
    # Include last 6 conversation turns so the model can ask clarifying questions.
    for turn in payload.history[-6:]:
        messages.append({"role": "user", "content": turn.get("user", "")})
        if turn.get("assistant"):
            messages.append({"role": "assistant", "content": turn["assistant"]})
    messages.append({"role": "user", "content": user_text})

    text = ""
    last_error = None

    # Try Groq first (free Llama), fall back to OpenAI
    if get_groq_api_key():
        try:
            groq_client = get_groq_client()
            # Discover the first available chat model on the account — Groq has
            # changed model names over time and "not found" on the right tier
            # is the most common reason a free key stops working.
            available_models = []
            try:
                models_response = await groq_client.models.list()
                # Accept any model that could be a chat model, including namespaced ones
                available_models = [
                    m.id for m in models_response.data
                    if any(tag in m.id.lower() for tag in [
                        "llama", "mixtral", "gemma", "qwen", "allam", "compound", "gpt-oss"
                    ])
                ]
            except Exception as e:
                logger.warning("Could not list Groq models: %s", e)

            # Prefer smaller/faster models first; all are free tier on Groq.
            # Models prefixed with "groq/", "openai/" etc. need the full ID.
            preferred = [
                "qwen/qwen3.8-27b",
                "qwen/qwen3.6-27b",
                "allam-2-7b",
                "groq/compound-mini",
                "groq/compound",
                "openai/gpt-oss-20b",
                "openai/gpt-oss-120b",
                "llama-3.1-8b-instant",
                "llama-3.3-70b-versatile",
            ]
            chosen_model = None
            for p in preferred:
                if p in available_models:
                    chosen_model = p
                    break
            if not chosen_model and available_models:
                chosen_model = available_models[0]

            if not chosen_model:
                # Fall through to a known-good name; the call will 404 if it's
                # truly gone, and the OpenAI fallback picks up after that.
                chosen_model = "llama-3.1-8b-instant"

            logger.info(f"Using Groq model: {chosen_model} (available: {available_models[:5]})")

            response = await groq_client.chat.completions.create(
                model=chosen_model,
                messages=messages,
                temperature=0.2,
            )
            text = response.choices[0].message.content or ""
        except Exception as e:
            logger.warning("Groq chat failed, falling back to OpenAI: %s", e)
            last_error = e

    # Fallback to OpenAI if Groq failed or not configured
    if not text and get_openai_api_key():
        try:
            openai_client = get_openai_client()
            response = await openai_client.chat.completions.create(
                model=os.environ.get("OPENAI_MODEL", "gpt-4o-mini"),
                messages=messages,
                temperature=0.2,
                response_format={"type": "json_object"},
            )
            text = response.choices[0].message.content or ""
        except Exception as e:
            logger.exception("OpenAI chat failed")
            last_error = e

    if not text:
        raise HTTPException(status_code=502, detail="The assistant is unavailable. Please add credits to your AI provider account.")

    cleaned = text.strip()
    if cleaned.startswith("```"):
        parts = cleaned.split("```")
        # A reply that opens a fence but never closes it yields a single part;
        # indexing [1] used to raise IndexError and turn it into a 500.
        cleaned = parts[1] if len(parts) > 1 else parts[0]
        if cleaned.startswith("json"):
            cleaned = cleaned[4:]
    start, end = cleaned.find("{"), cleaned.rfind("}")
    parsed = None
    if start != -1 and end > start:
        try:
            parsed = json.loads(cleaned[start:end + 1])
        except Exception:
            parsed = None

    if not isinstance(parsed, dict):
        return {"reply": _clean_str(text, MAX_REPLY_LEN), "actions": [], "transcript": transcript}

    return {
        "reply": _clean_str(parsed.get("reply"), MAX_REPLY_LEN),
        "actions": sanitize_actions(parsed.get("actions")),
        "transcript": transcript,
    }


# ---------------------------------------------------------------------------
# Data Import: PDF/DOCX parsing for onboarding migration
# ---------------------------------------------------------------------------
import re
import io
from datetime import datetime

ALLOWED_IMPORT_EXT = {".pdf", ".docx"}
MAX_IMPORT_BYTES = 10 * 1024 * 1024  # 10 MB


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def normalize_phone(phone_str: str) -> str:
    """Normalize phone number to last 10 digits."""
    if not phone_str:
        return ""
    digits = re.sub(r'\D', '', phone_str)
    # Take last 10 digits
    return digits[-10:] if len(digits) >= 10 else digits


def parse_amount(amount_str: str) -> Optional[float]:
    """Parse amount string to float, handling Indian formats."""
    if not amount_str:
        return None
    # Remove currency symbols and spaces
    cleaned = re.sub(r'[₹$€£¥]', '', amount_str)
    cleaned = cleaned.replace(',', '').strip()
    # Handle Hindi numerals
    hindi_digits = {'०': '0', '१': '1', '२': '2', '३': '3', '४': '4',
                    '५': '5', '६': '6', '७': '7', '८': '8', '९': '9'}
    for hindi, eng in hindi_digits.items():
        cleaned = cleaned.replace(hindi, eng)
    try:
        amount = float(cleaned)
        if amount > 0 and math.isfinite(amount):
            return round(amount, 2)
    except (ValueError, TypeError):
        pass
    return None


def parse_date(date_str: str) -> Optional[str]:
    """Parse date string to ISO format."""
    if not date_str:
        return None
    date_str = date_str.strip()

    # Common Indian formats
    formats = [
        '%d/%m/%Y', '%d-%m-%Y', '%d.%m.%Y',
        '%d/%m/%y', '%d-%m-%y', '%d.%m.%y',
        '%Y-%m-%d', '%d %b %Y', '%d %B %Y',
    ]

    for fmt in formats:
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.isoformat()
        except ValueError:
            continue
    return None


def parse_ledger_text(text: str) -> dict:
    """Parse extracted text to find parties, transactions, and opening balances.

    Looks for common Indian ledger patterns:
    - Names with phone numbers
    - Per-line transactions: date + amount + Gave/Got
    - Summary lines: Total Due / Balance / Net that set openingBalance
    """
    parties = {}  # name -> {phone, type, openingBalance}
    transactions = []
    warnings = []

    if not text or len(text.strip()) < 10:
        return {
            "success": False,
            "parties": [],
            "transactions": [],
            "warnings": ["Document appears to be empty or too short to contain ledger data."],
        }

    lines = text.split('\n')
    current_party = None

    # Pattern 1: Phone numbers (Indian 10-digit)
    phone_pattern = re.compile(r'(?:\+91[\s\-]?)?(\d{10})\b')

    # Pattern 2: Amounts (currency optional)
    amount_pattern = re.compile(r'(?:₹|Rs\.?|INR)?\s*([\d,]+(?:\.\d{1,2})?)')

    # Pattern 3: Dates
    date_patterns = [
        re.compile(r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b'),
        re.compile(r'\b(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{2,4})\b', re.IGNORECASE),
    ]

    # Pattern 4: Transaction type indicators
    gave_pattern = re.compile(r'\b(gave|given|udhar|udhaar|दिए|दिया|debit|dr\.?)\b', re.IGNORECASE)
    got_pattern = re.compile(r'\b(got|received|मिले|मिला|credit|cr\.?|paid|payment)\b', re.IGNORECASE)

    # Pattern 5: Name + Phone on same line
    name_phone_pattern = re.compile(
        r'([A-Z][a-zA-Z\s]{2,40})\s*[\-:|\s]+\s*(\+?91[\s\-]?)?(\d{10})'
    )

    # Pattern 6: Balance / Total Due summary lines (sets openingBalance for current party)
    # Matches patterns like "Total: ₹500", "Total Due: ₹500", "Balance: ₹500",
    # "Net: ₹500", "Due: ₹500", "₹500" alone on a line near the party header
    balance_summary_pattern = re.compile(
        r'(?:total|due|balance|net|pending|बकाया|₹|Rs\.?)?[\s:]*'
        r'(?:₹|Rs\.?)?\s*([\d,]+(?:\.\d{1,2})?)',
        re.IGNORECASE
    )
    balance_indicator_pattern = re.compile(
        r'\b(total\s*(?:due|amount)?|balance|net\s*(?:due|amount)?|pending|due|'
        r'बकाया|कुल\s*देय|कुल|शेष)\b',
        re.IGNORECASE
    )

    def _is_balance_line(line: str) -> bool:
        """Return True if this line looks like a balance/total summary."""
        stripped = line.strip()
        if not stripped:
            return False
        # Must contain an amount
        if not amount_pattern.search(stripped):
            return False
        # Must contain a balance indicator keyword
        if balance_indicator_pattern.search(stripped):
            return True
        # Also catch lines that are just "₹500" or "Rs 500" alone (short lines)
        if len(stripped) < 20 and re.match(r'^(?:₹|Rs\.?|INR)\s*[\d,]+(?:\.\d+)?$', stripped, re.IGNORECASE):
            return True
        return False

    for line in lines:
        line = line.strip()
        if not line or len(line) < 3:
            continue

        # ── 1. Name + Phone (party header) ───────────────────────────────────
        name_phone_match = name_phone_pattern.search(line)
        if name_phone_match:
            name = name_phone_match.group(1).strip()
            phone = normalize_phone(name_phone_match.group(3))
            if name and phone and name not in parties:
                parties[name] = {
                    "name": name,
                    "phone": phone,
                    "type": "CUSTOMER",
                    "openingBalance": 0,
                }
                current_party = name
                continue

        # ── 2. Just a phone number (associate with current party) ─────────────
        phone_match = phone_pattern.search(line)
        if phone_match and current_party:
            phone = normalize_phone(phone_match.group(1))
            if parties.get(current_party, {}).get("phone") in (None, ""):
                parties[current_party]["phone"] = phone

        # ── 3. Balance / Total Due line (sets openingBalance) ────────────────
        if current_party and _is_balance_line(line):
            # Extract the largest amount on this line as the balance
            amounts_in_line = amount_pattern.findall(line)
            if amounts_in_line:
                # Take the last/largest amount — total lines usually have one, but
                # fall back to the last one if there are multiple
                balance_amount = max(parse_amount(a) or 0 for a in amounts_in_line)
                if balance_amount > 0 and parties[current_party]["openingBalance"] == 0:
                    parties[current_party]["openingBalance"] = balance_amount

        # ── 4. Per-line transaction: date + amount + Gave/Got ────────────────
        amount_match = amount_pattern.search(line)
        if amount_match:
            amount = parse_amount(amount_match.group(1))
            if amount and amount > 0:
                is_gave = bool(gave_pattern.search(line))
                is_got = bool(got_pattern.search(line))

                # Only treat as transaction if it has a gave/got indicator
                if is_gave or is_got:
                    # Find date in this line
                    date_iso = None
                    for dp in date_patterns:
                        date_match = dp.search(line)
                        if date_match:
                            date_iso = parse_date(date_match.group(1))
                            break

                    party_name = current_party
                    if not party_name:
                        words = line.split()
                        if words and words[0][0].isupper() and len(words[0]) > 2:
                            # Skip lines that are just amounts (no meaningful name)
                            if not re.match(r'^[\d₹\s,.:\-]+$', line):
                                party_name = words[0]

                    if party_name:
                        transactions.append({
                            "partyName": party_name,
                            "amount": amount,
                            "type": "DEBIT" if is_gave else "CREDIT",
                            "note": "Imported from document",
                            "date": date_iso or datetime.now().isoformat(),
                        })

    # Convert parties dict to list
    parties_list = list(parties.values())

    # If no parties found but we have transactions, create parties from transaction names
    if not parties_list and transactions:
        seen_names = set()
        for tx in transactions:
            name = tx["partyName"]
            if name not in seen_names:
                seen_names.add(name)
                parties_list.append({
                    "name": name,
                    "phone": "",
                    "type": "CUSTOMER",
                    "openingBalance": 0,
                })

    if not parties_list and not transactions:
        warnings.append("No ledger data could be extracted. The document format may not be recognized.")
    elif len(transactions) == 0 and parties_list:
        warnings.append(f"Found {len(parties_list)} parties but no transactions. Please add transactions manually.")
    else:
        warnings.append(f"Extracted {len(parties_list)} parties and {len(transactions)} transactions. Please review before confirming.")

    return {
        "success": len(parties_list) > 0 or len(transactions) > 0,
        "parties": parties_list,
        "transactions": transactions,
        "warnings": warnings,
    }


@api_router.post("/import/parse")
async def import_parse(file: UploadFile = File(...), user: Optional[dict] = Depends(get_optional_user)):
    """Parse uploaded PDF or DOCX file to extract ledger data.

    Returns extracted parties and transactions for user review.

    Auth is optional: parsing a file does not need a user account, and the
    import flow runs during onboarding before the user has signed in. If a
    valid token is present we use it for rate limiting.
    """
    if user:
        enforce_user_rate_limit(str(user["user_id"]))

    suffix = Path(file.filename or "document").suffix.lower()
    if suffix not in ALLOWED_IMPORT_EXT:
        raise HTTPException(status_code=415, detail=f"Unsupported file format: {suffix}. Please upload PDF or DOCX files.")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Empty file")
    if len(content) > MAX_IMPORT_BYTES:
        raise HTTPException(status_code=413, detail="File must be 10 MB or smaller")

    # Extract text based on file type
    if suffix == ".pdf":
        text = extract_text_from_pdf(content)
    else:  # .docx
        text = extract_text_from_docx(content)

    if not text:
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from the file. The file may be corrupted, password-protected, or contain only images."
        )

    # Parse the extracted text
    result = parse_ledger_text(text)

    logger.info(
        f"Import parse for user {user['user_id']}: "
        f"{len(result['parties'])} parties, {len(result['transactions'])} transactions"
    )

    return result


# ── #21 SMS Auto-Parsing ──────────────────────────────────────────
class SmsParseRequest(BaseModel):
    sms_text: str = Field(..., max_length=2000)
    sender: Optional[str] = None


class SmsParseResponse(BaseModel):
    amount: Optional[float] = None
    party_name: Optional[str] = None
    party_phone: Optional[str] = None
    type: Optional[str] = None  # "DEBIT" (you paid) or "CREDIT" (you received)
    reference_id: Optional[str] = None
    date: Optional[str] = None
    confidence: float = 0.0
    raw_text: str = ""
    parser: str = "regex"


# Common Indian-bank UPI/card SMS patterns. Each tuple is (compiled regex, type)
# `type` is "DEBIT" if the user paid out, "CREDIT" if they received money.
_SMS_PATTERNS = [
    # HDFC: "Rs.500.00 debited from a/c **1234 to VPA merchant@okaxis on 01-09-2026"
    (re.compile(
        r"Rs\.?\s*([\d,]+(?:\.\d{1,2})?)\s*(?:has been\s*)?debited\s*(?:from\s*(?:a/c|account)\s*[\w\*]+)?\s*(?:to\s*(?:VPA\s*)?([\w.\-@]+))?",
        re.IGNORECASE), "DEBIT"),
    # "credited" / "received" patterns
    (re.compile(
        r"Rs\.?\s*([\d,]+(?:\.\d{1,2})?)\s*(?:has been\s*)?credited\s*(?:to\s*(?:a/c|account)\s*[\w\*]+)?\s*(?:from\s*(?:VPA\s*)?([\w.\-@]+))?",
        re.IGNORECASE), "CREDIT"),
    # Generic spent/sent: "You spent Rs 500 at Swiggy"
    (re.compile(
        r"(?:spent|sent|paid)\s*Rs\.?\s*([\d,]+(?:\.\d{1,2})?)\s*(?:to|at|on)?\s*([A-Za-z][A-Za-z0-9 _\-]{1,40})",
        re.IGNORECASE), "DEBIT"),
    # Generic received: "You received Rs 500 from John"
    (re.compile(
        r"(?:received|got)\s*Rs\.?\s*([\d,]+(?:\.\d{1,2})?)\s*(?:from|by)?\s*([A-Za-z][A-Za-z0-9 _\-]{1,40})",
        re.IGNORECASE), "CREDIT"),
    # UPI txn ID: "UPI Ref 1234567890" or "Txn ID ABC123"
    (re.compile(r"(?:UPI\s*Ref|Txn\s*ID|UPI\s*Txn\s*ID|Reference)[:\s]+([A-Z0-9]{8,20})", re.IGNORECASE), None),
    # Date patterns: 01-09-2026 or 01/09/26
    (re.compile(r"(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})"), None),
]
_AMOUNT_RE = re.compile(r"Rs\.?\s*([\d,]+(?:\.\d{1,2})?)", re.IGNORECASE)
_PHONE_RE = re.compile(r"(\+?91[\s\-]?)?(\d{10})")


def _parse_sms_regex(text: str) -> dict:
    """Try to extract amount, party, type, ref id, date from common SMS patterns.
    Returns a dict with whatever was found. Confidence is a 0-1 heuristic.
    """
    result: dict = {"amount": None, "party": None, "type": None, "ref": None, "date": None}
    confidence = 0.0
    for pat, t in _SMS_PATTERNS:
        m = pat.search(text)
        if not m:
            continue
        if t in ("DEBIT", "CREDIT") and result["type"] is None:
            result["type"] = t
            if t == "CREDIT" and not result["party"]:
                # group 2 = "from VPA/name"
                candidate = m.group(2) if m.lastindex and m.lastindex >= 2 else None
                if candidate:
                    result["party"] = candidate.split("@")[0] if "@" in candidate else candidate
                    confidence += 0.4
            elif t == "DEBIT" and not result["party"]:
                candidate = m.group(2) if m.lastindex and m.lastindex >= 2 else None
                if candidate:
                    result["party"] = candidate.split("@")[0] if "@" in candidate else candidate
                    confidence += 0.4
        elif t is None and "Ref|Txn" in pat.pattern and result["ref"] is None:
            result["ref"] = m.group(1)
            confidence += 0.1
        elif t is None and "date" in pat.pattern.lower() and result["date"] is None:
            result["date"] = m.group(1)
    # Always try to pull the first rupee amount out
    if not result["amount"]:
        am = _AMOUNT_RE.search(text)
        if am:
            try:
                result["amount"] = float(am.group(1).replace(",", ""))
                confidence += 0.3
            except ValueError:
                pass
    # Look for a phone number
    pm = _PHONE_RE.search(text)
    if pm:
        result["phone"] = pm.group(2)
        confidence += 0.1
    return {**result, "confidence": min(1.0, confidence)}


async def _parse_sms_llm(text: str) -> dict:
    """Fallback: ask Groq to extract structured fields from a free-form SMS body."""
    groq_client = get_groq_client()
    prompt = (
        "Extract a transaction from this Indian bank UPI/debit/credit SMS. "
        "Return ONLY valid JSON with keys: amount (number), party_name (string), "
        "type ('DEBIT' or 'CREDIT'), reference_id (string), date (DD-MM-YYYY). "
        "If a field is missing, use null.\n\n"
        f"SMS:\n{text}"
    )
    response = await groq_client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[
            {"role": "system", "content": "You are an SMS-to-JSON extractor. Reply with only JSON."},
            {"role": "user", "content": prompt},
        ],
        temperature=0,
        max_tokens=200,
        response_format={"type": "json_object"},
    )
    raw = response.choices[0].message.content or "{}"
    try:
        parsed = json.loads(raw)
        parsed.setdefault("amount", None)
        parsed.setdefault("party_name", None)
        parsed.setdefault("type", None)
        parsed.setdefault("reference_id", None)
        parsed.setdefault("date", None)
        return parsed
    except json.JSONDecodeError:
        return {"amount": None, "party_name": None, "type": None, "reference_id": None, "date": None}


@api_router.post("/sms/parse", response_model=SmsParseResponse)
async def sms_parse(req: SmsParseRequest, user: dict = Depends(get_authenticated_user)):
    """Parse a bank UPI/credit/debit SMS and return a structured transaction.

    Tries a regex pass first; falls back to Groq if confidence is low.
    """
    enforce_user_rate_limit(str(user["user_id"]))
    text = (req.sms_text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="SMS text is empty")

    regex_result = _parse_sms_regex(text)
    if regex_result["confidence"] >= 0.6:
        return SmsParseResponse(
            amount=regex_result.get("amount"),
            party_name=regex_result.get("party"),
            party_phone=regex_result.get("phone"),
            type=regex_result.get("type"),
            reference_id=regex_result.get("ref"),
            date=regex_result.get("date"),
            confidence=regex_result["confidence"],
            raw_text=text,
            parser="regex",
        )

    # LLM fallback
    try:
        llm_result = await _parse_sms_llm(text)
        return SmsParseResponse(
            amount=llm_result.get("amount"),
            party_name=llm_result.get("party_name"),
            party_phone=regex_result.get("phone"),
            type=llm_result.get("type"),
            reference_id=llm_result.get("reference_id"),
            date=llm_result.get("date"),
            confidence=max(0.6, regex_result["confidence"] + 0.3),
            raw_text=text,
            parser="llm",
        )
    except HTTPException:
        # No API key — return the regex result as-is
        return SmsParseResponse(
            amount=regex_result.get("amount"),
            party_name=regex_result.get("party"),
            party_phone=regex_result.get("phone"),
            type=regex_result.get("type"),
            reference_id=regex_result.get("ref"),
            date=regex_result.get("date"),
            confidence=regex_result["confidence"],
            raw_text=text,
            parser="regex",
        )

# Include the router in the main app
app.include_router(api_router)

DEFAULT_ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "http://localhost:8081",
    "http://127.0.0.1:8081",
    "exp://localhost:8081",
    "exp://127.0.0.1:8081",
]

allowed_origins = [
    origin.strip()
    for origin in os.getenv("ALLOWED_ORIGINS", "").split(",")
    if origin.strip()
]
if not allowed_origins:
    allowed_origins = DEFAULT_ALLOWED_ORIGINS

app.add_middleware(
    CORSMiddleware,
    allow_credentials=False,
    allow_origins=allowed_origins,
    allow_methods=["GET", "POST", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)
